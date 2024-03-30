from crewai_tools import BaseTool
import docx
from docx.shared import Pt

class DocxReportGenerator(BaseTool):
    name: str = "DocxReportGenerator"
    description: str = "Generates a detailed comparison report on multi-agent frameworks in Word format."

    def _run(self, analysis_results: dict) -> str:
        # Create a new Word document
        doc = docx.Document()
        doc.add_heading('Multi-Agent Frameworks Comparison Report', 0)

        # Introduction
        doc.add_paragraph(
            "This document presents a comprehensive comparison of various multi-agent frameworks, "
            "evaluating them based on a set of criteria to inform decision-making."
        )

        # Generating a section for each framework
        for framework in analysis_results.get('frameworks', []):
            doc.add_heading(framework['name'], level=1)
            p = doc.add_paragraph()
            runner = p.add_run("Description: ")
            runner.bold = True
            p.add_run(framework['description'])

            # Criteria and Ratings
            for criterion, rating in framework.get('ratings', {}).items():
                p = doc.add_paragraph(style='ListBullet')
                p.add_run(f"{criterion}: {rating} stars")

        # Comparison Summary and Table
        doc.add_heading('Comparison Summary', level=1)
        doc.add_paragraph(analysis_results.get('summary', 'Summary of the comparison.'))

        # Adding a table for the comparison
        table = doc.add_table(rows=1, cols=len(analysis_results.get('criteria', [])) + 1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Framework'
        for i, criterion in enumerate(analysis_results.get('criteria', []), start=1):
            hdr_cells[i].text = criterion

        for framework in analysis_results.get('frameworks', []):
            row_cells = table.add_row().cells
            row_cells[0].text = framework['name']
            for i, criterion in enumerate(analysis_results.get('criteria', []), start=1):
                row_cells[i].text = str(framework['ratings'].get(criterion, 'N/A'))

        # Save the document
        report_path = "comparison_report.docx"
        doc.save(report_path)
        return f"Report generated at {report_path}"
